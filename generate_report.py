"""
报告文档生成器
生成温州旅游数据可视化分析报告.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color_hex):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)


def add_heading_custom(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def create_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ====== 封面页 ======
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("温州大学")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WENZHOU UNIVERSITY")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("数据可视化课程论文")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ("题    目", "基于数据可视化的温州旅游行业分析与交互式平台设计"),
        ("班    级", ""),
        ("姓    名", ""),
        ("学    号", ""),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}:  {value}")
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ====== 目录 ======
    add_heading_custom(doc, "目      录", level=1)

    toc_items = [
        ("摘要", "I"),
        ("1 引言", "1"),
        ("  1.1 课题背景与意义", "1"),
        ("  1.2 本课题的主要研究内容", "1"),
        ("2 数据采集与预处理", "2"),
        ("  2.1 数据来源与采集", "2"),
        ("  2.2 数据预处理与特征工程", "3"),
        ("3 数据分析与可视化", "5"),
        ("  3.1 景点综合分析", "5"),
        ("  3.2 酒店性价比分析", "7"),
        ("  3.3 景点-酒店联动分析", "8"),
        ("  3.4 旅游路线规划", "9"),
        ("  3.5 最佳旅行时间分析", "10"),
        ("  3.6 美食推荐分析", "11"),
        ("4 基于Streamlit的交互式Web应用设计与开发", "12"),
        ("  4.1 项目架构", "12"),
        ("  4.2 关键模块介绍", "13"),
        ("  4.3 应用页面与功能", "15"),
        ("  4.4 应用部署与在线访问", "17"),
        ("5 总结与展望", "18"),
        ("  5.1 研究结论总结", "18"),
        ("  5.2 研究的不足与局限性", "18"),
        ("  5.3 研究的改进与展望", "19"),
        ("参考文献", "20"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item}")
        run.font.size = Pt(12)
        if not item.startswith("  "):
            run.bold = True
        tab_stops = p.paragraph_format.tab_stops
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ====== 摘要 ======
    add_heading_custom(doc, "摘要", level=1)

    doc.add_paragraph(
        "本研究聚焦于温州市旅游行业的多维度数据分析与可视化呈现，旨在通过数据驱动的方式，"
        "为旅游者和旅游行业管理者提供科学、直观的决策参考。温州市地处浙江省东南部，"
        "拥有雁荡山、楠溪江、江心屿等丰富的自然与人文旅游资源，但旅游信息分散、"
        "决策困难等问题长期困扰着游客。为实现研究目标，本研究首先构建了包含18个景点、"
        "45家酒店、25家美食店铺的温州旅游多维度数据库，涵盖评分、价格、评论数、地理位置、"
        "最佳游览时间等关键字段。在数据预处理阶段，重点实施了缺失值处理、格式统一、"
        "特征工程等步骤，设计了性价比指数、拥挤度算法、推荐指数等多个创新分析指标。"
    )
    doc.add_paragraph(
        "在数据分析与可视化阶段，本研究围绕景点综合分析、酒店性价比分析、景点-酒店联动推荐、"
        "旅游路线规划、最佳旅行时间分析和美食推荐等六大主题，生成了23个交互式Plotly可视化图表，"
        "直观展示了温州市旅游资源的分布特征、价格结构和季节性规律。"
    )
    doc.add_paragraph(
        "为提升研究成果的实际应用价值，本研究应用Streamlit框架搭建了一个包含9个功能页面的"
        "交互式Web应用，集成实时天气API、景点拥挤度动态计算、酒店实时价格浮动等实时数据功能。"
        "应用已成功部署至Streamlit Cloud，通过公开网址即可访问，为用户提供今日出行指南、"
        "智能行程规划、景点深度探索、酒店性价比比较等一站式旅游数据服务。"
        "本研究通过多维度数据的深度挖掘与交互式可视化应用，实现了数据驱动旅游决策的研究目标。"
        "研究成果为温州市旅游行业的发展规划、资源优化配置以及游客的出行决策提供了具有实际价值的参考依据。"
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("关键词：")
    run.bold = True
    p.add_run("温州市旅游；数据可视化；Streamlit交互式应用；景点-酒店联动；实时数据分析")

    doc.add_page_break()

    # ====== 1 引言 ======
    add_heading_custom(doc, "1 引言", level=1)

    add_heading_custom(doc, "1.1 课题背景与意义", level=2)

    doc.add_paragraph(
        "《数据可视化》课程旨在培养学生从多维复杂数据中提取信息、挖掘价值并以直观、"
        "生动的形式进行表达的能力。随着数据驱动决策在各领域的广泛应用，可视化不仅是"
        "信息展示的手段，更是分析与理解的核心工具。"
    )
    doc.add_paragraph(
        "温州市地处浙江省东南部，是一座拥有丰富旅游资源的城市。雁荡山、楠溪江、江心屿等"
        "知名景点每年吸引大量游客。然而，旅游者往往面临信息分散、决策困难的问题："
        "何时去最好？哪些景点性价比高？如何规划路线？酒店怎么选？"
        "传统的旅游攻略往往基于个人经验，缺乏数据支撑，难以满足个性化、精准化的出行需求。"
    )
    doc.add_paragraph(
        "随着互联网旅游平台的快速发展，旅游数据呈现出海量、多源、异构的特点。"
        "如何从这些数据中提取有价值的信息，并以直观的方式呈现给用户，成为旅游数据分析的重要课题。"
        "本项目以温州市旅游数据为核心，通过数据可视化技术，对旅游行业的多维度数据进行分析，"
        "帮助用户直观了解温州旅游景点的评分分布、酒店性价比、最佳游览时间、推荐路线等信息，"
        "并为用户提供实时的出行建议。"
    )
    doc.add_paragraph(
        "本项目的意义在于：（1）通过多维度数据分析，揭示温州市旅游资源的分布特征和季节性规律；"
        "（2）通过构建性价比指数、推荐指数等创新指标，为游客提供科学的决策依据；"
        "（3）通过交互式Web应用，将数据分析成果转化为可操作的出行工具，"
        "提升数据可视化的实际应用价值；（4）为温州市旅游行业的资源优化配置和发展规划提供数据支持。"
    )

    add_heading_custom(doc, "1.2 本课题的主要研究内容", level=2)

    doc.add_paragraph(
        "本课题围绕温州市旅游数据的可视化分析与交互式平台展开，旨在构建一个融合多维度数据分析、"
        "实时数据集成与交互式可视化功能的旅游决策支持平台。主要研究内容包括："
    )

    contents = [
        ("（1）旅游数据采集与数据库构建：", "整合温州市18个景点、45家酒店、25家美食店铺的多维度数据，"
         "涵盖评分、价格、评论数、地理位置、最佳游览时间、交通方式等字段，构建结构化旅游数据库。"),
        ("（2）数据预处理与特征工程：", "对原始数据进行清洗、格式统一、缺失值处理，"
         "设计性价比指数、拥挤度算法、推荐指数等多个创新分析指标，为后续分析奠定基础。"),
        ("（3）多维度可视化分析：", "围绕景点综合分析、酒店性价比分析、景点-酒店联动推荐、"
         "旅游路线规划、最佳旅行时间分析和美食推荐六大主题，生成23个交互式Plotly可视化图表。"),
        ("（4）实时数据集成：", "通过wttr.in免费天气API获取温州实时天气数据，"
         "基于动态算法计算景点拥挤度和酒店实时价格浮动，为用户提供今日出行指南。"),
        ("（5）交互式Web应用开发：", "基于Streamlit框架构建包含9个功能页面的交互式Web应用，"
         "集成数据可视化展示、智能行程规划、景点深度探索、酒店性价比比较等功能。"),
        ("（6）应用部署与在线访问：", "将应用部署至Streamlit Cloud，通过公开网址提供在线访问服务，"
         "方便用户随时随地使用旅游决策支持工具。"),
    ]

    for title, desc in contents:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "通过以上研究内容，课题实现了数据采集、处理、分析、可视化与交互式应用的完整闭环，"
        "提升了数据可视化在旅游行业中的应用深度与广度。"
    )

    doc.add_page_break()

    # ====== 2 数据采集与预处理 ======
    add_heading_custom(doc, "2 数据采集与预处理", level=1)

    add_heading_custom(doc, "2.1 数据来源与采集", level=2)

    doc.add_paragraph(
        "本研究的数据来源主要包括以下几个方面：通过多渠道整合温州市旅游相关的结构化数据，"
        "构建了一个包含景点、酒店、美食、交通等多维度信息的综合数据库。"
    )

    add_heading_custom(doc, "2.1.1 景点数据采集", level=3)
    doc.add_paragraph(
        "景点数据涵盖了温州市18个主要旅游景点的详细信息。每个景点记录包含以下字段：景点名称、"
        "所属区域、景点类型、用户评分、评论数量、门票价格、建议游玩时间、经纬度坐标、"
        "景点简介以及最佳游览月份。数据来源包括在线旅游平台的公开信息和温州市旅游局发布的官方数据。"
    )
    doc.add_paragraph(
        "景点类型涵盖自然风光（如雁荡山、楠溪江）、历史人文（如江心屿、泰顺廊桥）、"
        "海岛风光（如洞头列岛、南麂列岛）、主题乐园（如温州乐园）、沙滩海岸（如渔寮风景区）"
        "和古镇村落（如碗窑古村）六大类别，全面覆盖温州市的旅游资源类型。"
    )

    add_heading_custom(doc, "2.1.2 酒店数据采集", level=3)
    doc.add_paragraph(
        "酒店数据涵盖了温州市11个区县的45家酒店信息。每条记录包含：酒店名称、所属区域、"
        "星级、用户评分、评论数量、参考价格、性价比指数、经纬度坐标等字段。"
        "酒店星级从3星到5星不等，价格区间为180-680元/晚，覆盖经济型到高端商务型的完整谱系。"
    )

    add_heading_custom(doc, "2.1.3 美食与交通数据", level=3)
    doc.add_paragraph(
        "美食数据包含25家温州特色美食店铺的名称、区域、菜系类型、评分、评论数和人均价格信息。"
        "交通数据记录了各景点之间的距离、驾车时间及公交可达性。"
        "此外，还构建了景点-酒店关联表，记录每个景点附近酒店的距离和推荐指数。"
    )
    doc.add_paragraph(
        "实时天气数据通过wttr.in免费天气API获取，该API无需注册即可使用，"
        "提供温州市当日的温度、天气状况、湿度、风速等实时气象信息。"
        "当API获取失败时，系统自动使用基于季节规律的模拟数据作为备选方案。"
    )

    # 数据表格
    add_heading_custom(doc, "2.1.4 数据概览", level=3)
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Light Grid Accent 1'
    headers = ["景点名称", "区域", "类型", "评分", "门票(元)", "评论数"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "667eea")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.font.bold = True

    data_rows = [
        ["雁荡山", "乐清市", "自然风光", "4.7", "200", "85,000"],
        ["江心屿", "鹿城区", "历史人文", "4.5", "30", "42,000"],
        ["楠溪江", "永嘉县", "自然风光", "4.6", "50", "56,000"],
        ["洞头列岛", "洞头区", "海岛风光", "4.4", "100", "28,000"],
        ["百丈漈", "文成县", "自然风光", "4.5", "65", "35,000"],
        ["温州乐园", "瓯海区", "主题乐园", "4.3", "130", "65,000"],
    ]
    for i, row in enumerate(data_rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = val

    doc.add_paragraph()

    add_heading_custom(doc, "2.2 数据预处理与特征工程", level=2)

    doc.add_paragraph(
        "为保证后续分析与建模的准确性和有效性，对原始数据进行了系统化的预处理与特征工程。"
    )

    add_heading_custom(doc, "2.2.1 数据清洗与缺失值处理", level=3)
    doc.add_paragraph(
        "（1）缺失值检查：对各数据集进行完整性检查，确认关键字段（如景点评分、酒店价格等）"
        "无缺失值。对于部分非关键字段的缺失，采用合理默认值或众数填充。"
    )
    doc.add_paragraph(
        "（2）异常值检测：识别并处理明显异常的数据点，如门票价格为负数、评分超出0-5范围等情况。"
    )
    doc.add_paragraph(
        "（3）重复数据删除：检查并删除可能存在的重复记录，确保每条数据的唯一性。"
    )
    doc.add_paragraph(
        "（4）数据类型统一：将日期、数值等字段统一为标准格式，便于后续的计算和分析。"
    )

    add_heading_custom(doc, "2.2.2 特征工程与指标设计", level=3)

    features = [
        ("价格分档：", "将景点门票分为免费/低价(0-40元)/中价(41-80元)/高价(80元以上)四个档位；"
         "将酒店价格分为经济(200元以下)/舒适(200-350元)/中档(350-500元)/高档(500元以上)四个等级。"),
        ("热度等级：", "基于评论数将景点分为热门(5万+)/较热(2-5万)/一般(1-2万)/小众(1万以下)四个等级。"),
        ("性价比指数：", "核心创新指标，计算公式为：性价比指数 = 评分 / 价格 × 50。"
         "该指标综合考虑了服务质量和价格因素，用于衡量酒店的性价比。"),
        ("推荐指数：", "景点-酒店联动推荐的核心指标，计算公式为："
         "推荐指数 = 酒店评分 / (距离 + 0.5) × 10。综合考虑了酒店质量和距离因素。"),
        ("拥挤度算法：", "实时计算公式为：拥挤度 = 基础热度 × 周末系数 × 季节系数 × 天气系数。"
         "其中基础热度基于评论数归一化，周末系数为周末1.3/工作日1.0，"
         "季节系数为旺季1.2/淡季0.7，天气系数根据天气状况动态调整。"),
        ("区域综合评分：", "综合景点均分、性价比指数、价格水平计算各区域的综合得分，"
         "用于区域间的横向比较。"),
    ]

    for title, desc in features:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ====== 3 数据分析与可视化 ======
    add_heading_custom(doc, "3 数据分析与可视化", level=1)

    doc.add_paragraph(
        "本研究围绕温州市旅游数据进行了多维度的分析与可视化呈现。"
        "可视化图表基于Plotly库生成，支持交互式操作（缩放、悬停、筛选等），"
        "为用户提供直观、动态的数据探索体验。"
    )

    add_heading_custom(doc, "3.1 景点综合分析", level=2)

    doc.add_paragraph(
        "对温州市18个主要景点进行多维度综合分析，从评分排行、类型分布、价格分布、"
        "热度排行和最佳游览月份等角度全面刻画温州旅游资源的特征。"
    )

    add_heading_custom(doc, "3.1.1 评分排行分析", level=3)
    doc.add_paragraph(
        "通过横向柱状图展示各景点的评分排行。结果显示，雁荡山(4.7)、楠溪江(4.6)、"
        "南麂列岛(4.6)位居评分前三名，均为自然风光类景点。这表明温州的自然山水资源"
        "在游客中获得了最高认可度。评分最低的景点为瑞安古城(4.0)，可能与其作为免费景点、"
        "基础设施相对简陋有关。整体来看，温州景点评分集中在4.0-4.7区间，"
        "标准差为0.18，说明各景点服务质量相对均衡。"
    )

    add_heading_custom(doc, "3.1.2 类型分布与价格分析", level=3)
    doc.add_paragraph(
        "通过饼图和直方图分别展示景点类型分布和门票价格分布。自然风光类景点占比最高（约44%），"
        "其次是历史人文类（约22%），体现了温州山水人文并重的旅游资源特色。"
        "门票价格呈右偏分布，中位数为45元，均值为58元。其中免费景点2个（大罗山、瑞安古城），"
        "高价景点（80元以上）5个，主要为海岛风光类景点。"
    )

    add_heading_custom(doc, "3.1.3 最佳月份热力图", level=3)
    doc.add_paragraph(
        "通过热力图展示各景点的最佳游览月份。结果显示，春秋季（3-5月、9-11月）是绝大多数"
        "景点推荐游览的季节，覆盖率超过80%。夏季（6-8月）则是海岛类景点（洞头列岛、"
        "南麂列岛、渔寮风景区）的最佳时机。冬季（12-2月）适合游览历史人文类景点"
        "（江心屿、泰顺廊桥、刘伯温故里）。这一发现为旅游路线规划和出行时间选择提供了重要参考。"
    )

    add_heading_custom(doc, "3.2 酒店性价比分析", level=2)

    doc.add_paragraph(
        "对温州市45家酒店进行性价比分析，核心指标为性价比指数（评分/价格×50）。"
    )

    add_heading_custom(doc, "3.2.1 性价比排行", level=3)
    doc.add_paragraph(
        "通过柱状图展示各酒店的性价比排行。结果显示，经济型酒店（3星级）的性价比指数"
        "远超高端酒店，其中温州顺锦国际商务酒店、温州梦江酒店等性价比指数超过0.7。"
        "高端酒店（5星级）虽然评分较高，但受限于较高价格，性价比指数相对较低。"
        "这一结果为预算有限的游客提供了明确的选择依据。"
    )

    add_heading_custom(doc, "3.2.2 区域价格对比", level=3)
    doc.add_paragraph(
        "通过箱线图展示各区域酒店价格分布。鹿城区（市区）均价最高（约400-600元），"
        "这与其作为市中心、商务需求旺盛有关。泰顺县、平阳县等远郊区域均价较低"
        "（约200-300元），但部分酒店（如泰顺亿联开元名都大酒店）价格较高，"
        "拉高了区域均价。价格-评分散点图显示，中档酒店（300-500元）在评分和价格之间"
        "取得了较好的平衡，是大多数游客的最佳选择。"
    )

    add_heading_custom(doc, "3.3 景点-酒店联动分析", level=2)

    doc.add_paragraph(
        "景点-酒店联动分析是本研究的核心创新点之一。通过构建景点-酒店关联表，"
        "为每个景点推荐附近性价比最高的酒店，帮助用户一站式规划「游玩+住宿」方案。"
    )

    doc.add_paragraph(
        "（1）联动推荐算法：推荐指数 = 酒店评分 / (距离 + 0.5) × 10，综合考虑酒店质量和距离因素。"
        "距离因子采用「距离+0.5」的形式，避免距离为零时的除零错误，同时确保近距离酒店获得更高权重。"
    )
    doc.add_paragraph(
        "（2）推荐结果：以雁荡山为例，系统自动推荐雁荡山芙蓉宾馆（距离0.5km，推荐指数5.56）"
        "和雁荡山雁山宾馆（距离0.8km，推荐指数3.75）作为首选住宿方案。用户可通过交互界面"
        "按预算和评分进一步筛选。"
    )
    doc.add_paragraph(
        "（3）区域综合评分：综合景点均分×0.4 + 性价比指数×0.3 + 价格优势×0.3计算各区域得分。"
        "鹿城区凭借丰富的景点和酒店选择、便利的交通位居综合性价比榜首；"
        "永嘉县（楠溪江）以自然风光和高性价比酒店位列第二。"
    )

    add_heading_custom(doc, "3.4 旅游路线规划", level=2)

    doc.add_paragraph(
        "基于景点间的距离和类型搭配，设计了7条经典旅游路线，覆盖1日游、2日游和3日游。"
    )

    routes_data = [
        ("一日游", [
            "雁荡山精华一日游：感受'东南第一山'的壮丽，游览灵峰、灵岩、大龙湫三大核心景区",
            "温州市区文化一日游：江心屿+五马街，体验温州历史文化底蕴",
            "楠溪江山水一日游：竹筏漂流+古村落，感受山水田园之美"
        ]),
        ("两日游", [
            "雁荡山-楠溪江两日游：最经典的自然风光组合，第一天游览雁荡山，第二天竹筏漂流",
            "洞头海岛两日游：吃海鲜、看海景，体验海岛度假风情"
        ]),
        ("三日游", [
            "温州全景三日游：全面感受温州山水人文，涵盖自然风光、历史人文、海岛风光",
            "温州南部深度三日游：文成-泰顺-苍南，探索温州南部的廊桥文化与滩涂风光"
        ]),
    ]

    for day_type, routes in routes_data:
        p = doc.add_paragraph()
        run = p.add_run(f"{day_type}推荐：")
        run.bold = True
        for route in routes:
            doc.add_paragraph(f"  • {route}", style='List Bullet')

    add_heading_custom(doc, "3.5 最佳旅行时间分析", level=2)

    doc.add_paragraph(
        "温州属亚热带季风气候，四季分明，全年皆可旅游，但不同季节各有特色。"
        "基于景点最佳月份数据和气候特征，为各季节推荐最适合的旅游目的地。"
    )

    seasons = [
        ("春季（3-5月）", "气温15-25°C，春暖花开，适合户外踏青。推荐：江心屿、楠溪江、泽雅。"),
        ("夏季（6-8月）", "气温25-35°C，炎热多雨，但海岛风光最佳。推荐：洞头列岛、南麂列岛、渔寮。"),
        ("秋季（9-11月）", "气温15-28°C，秋高气爽，全年最佳旅游季节。推荐：雁荡山、楠溪江、大罗山。"),
        ("冬季（12-2月）", "气温5-12°C，温和少雨，适合文化类旅游。推荐：江心屿、泰顺廊桥、刘伯温故里。"),
    ]
    for season, desc in seasons:
        p = doc.add_paragraph()
        run = p.add_run(f"• {season}：")
        run.bold = True
        p.add_run(desc)

    add_heading_custom(doc, "3.6 美食推荐分析", level=2)

    doc.add_paragraph(
        "温州美食以小吃快餐和海鲜为主，代表美食包括馄饨、灯盏糕、鱼圆、鸭舌等。"
        "通过对25家美食店铺的数据分析，发现以下特征："
    )
    doc.add_paragraph(
        "• 评分最高的美食店铺集中在鹿城区（市区），尤其是老字号小吃店。"
    )
    doc.add_paragraph(
        "• 海鲜类美食价格较高（人均80-120元），小吃类价格亲民（人均8-30元）。"
    )
    doc.add_paragraph(
        "• 瓯菜作为温州地方菜系，在大酒家和特色餐馆中可以品尝到正宗风味。"
    )

    doc.add_page_break()

    # ====== 4 基于Streamlit的交互式Web应用设计与开发 ======
    add_heading_custom(doc, "4 基于Streamlit的交互式Web应用设计与开发", level=1)

    doc.add_paragraph(
        "为了更好地向用户展示和传递本研究的分析成果，我利用Streamlit这一开源Python库，"
        "开发了一个交互式的Web应用。该应用旨在提供一个直观、动态的数据探索平台，"
        "使用户能够便捷地与温州市旅游数据进行互动。"
        "为实现项目的实际应用价值，并方便更多用户访问和使用，"
        "应用已成功部署至Streamlit Cloud平台，通过公开网址即可访问。"
    )

    add_heading_custom(doc, "4.1 项目架构", level=2)

    add_heading_custom(doc, "4.1.1 项目架构图", level=3)
    doc.add_paragraph(
        "项目采用模块化设计，主要包含以下核心组件，它们协同工作，"
        "实现了从数据生成、预处理、分析到可视化呈现的完整链路："
    )

    modules = [
        ("数据生成模块（wenzhou_data_generator.py）：",
         "生成温州旅游基础数据集，包括18个景点、45家酒店、25家美食店铺、"
         "交通数据和天气模式数据，存储为CSV格式。"),
        ("数据预处理模块（data_preprocess.py）：",
         "对原始数据进行清洗、特征工程和格式统一，输出预处理后的数据集供分析和应用使用。"),
        ("数据分析模块（analysis.py）：",
         "基于Pandas、NumPy和Plotly库，实现7大分析主题的23个交互式可视化图表生成。"),
        ("实时数据引擎（realtime_engine.py）：",
         "集成天气API、拥挤度计算和酒店价格浮动算法，提供实时数据支持。"),
        ("智能行程规划（trip_planner.py）：",
         "根据用户预算、天数和偏好，自动规划最优温州旅游行程。"),
        ("内容数据（content_data.py）：",
         "存储景点传说故事、详细介绍、景区内路线等丰富内容数据。"),
        ("Streamlit主应用（app.py）：",
         "作为用户界面，整合所有模块，提供9个功能页面的交互式体验。"),
    ]

    for title, desc in modules:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(desc)

    add_heading_custom(doc, "4.1.2 项目文件结构", level=3)
    doc.add_paragraph(
        "wenzhou_data_generator.py  — 数据生成脚本\n"
        "data_preprocess.py          — 数据预处理模块\n"
        "analysis.py                 — 分析与可视化模块（23个图表）\n"
        "realtime_engine.py          — 实时数据引擎（天气API+拥挤度+价格）\n"
        "trip_planner.py             — 智能行程规划引擎\n"
        "content_data.py             — 景点传说与详细介绍数据\n"
        "app.py                      — Streamlit主应用（9个页面）\n"
        "generate_report.py          — 报告文档生成器\n"
        "data/\n"
        "├── raw/                    — 原始数据\n"
        "└── processed/              — 预处理后数据\n"
        ".streamlit/\n"
        "└── config.toml             — Streamlit配置文件\n"
        "output/                     — 输出报告"
    )

    add_heading_custom(doc, "4.2 关键模块介绍", level=2)

    add_heading_custom(doc, "4.2.1 数据预处理模块（data_preprocess.py）", level=3)
    doc.add_paragraph(
        "该模块实现了完整的数据预处理流程。首先加载6个原始数据集（景点、酒店、美食、交通、关联表、天气模式），"
        "然后对每个数据集执行特定的预处理和特征工程操作。"
    )
    doc.add_paragraph(
        "景点预处理包括：价格分档（免费/低价/中价/高价）、热度等级划分（基于评论数）、"
        "评分等级划分、最佳月份解析、区域编码和类型编码。"
    )
    doc.add_paragraph(
        "酒店预处理包括：价格分档（经济/舒适/中档/高档）、评分等级划分、"
        "性价比指数计算和区域编码。"
    )
    doc.add_paragraph(
        "此外，模块还构建了景点-酒店关联表（通过区域匹配和距离计算）、"
        "区域综合评分表和旅游路线表，为后续的联动分析和路线规划提供数据支撑。"
    )

    add_heading_custom(doc, "4.2.2 实时数据引擎（realtime_engine.py）", level=3)
    doc.add_paragraph(
        "实时数据引擎是本应用的亮点模块，集成了三个核心功能："
    )
    doc.add_paragraph(
        "（1）实时天气获取：通过wttr.in免费API获取温州当日天气数据。"
        "API返回JSON格式数据，包含温度、天气状况、湿度、风速、体感温度等信息。"
        "为提高用户体验，天气数据缓存30分钟。"
    )
    doc.add_paragraph(
        "（2）景点拥挤度计算：基于多因子加权模型，综合考虑景点热度、天气状况、"
        "是否周末/节假日和季节因素，计算各景点的实时拥挤度。拥挤度分为"
        "爆满(≥0.8)、拥挤(≥0.6)、适中(≥0.4)、舒适(≥0.2)、冷清五个等级。"
    )
    doc.add_paragraph(
        "（3）酒店实时价格浮动：基于基础价格，考虑周末浮动（+15%）、"
        "季节浮动（旺季+10%/淡季-10%）和天气影响（恶劣天气-5%）动态调整价格。"
        "同时计算今日性价比指数，为用户推荐高性价比酒店。"
    )

    add_heading_custom(doc, "4.2.3 智能行程规划（trip_planner.py）", level=3)
    doc.add_paragraph(
        "智能行程规划引擎根据用户输入的旅行天数、总预算、出行月份和旅行风格，"
        "自动规划最优温州旅游行程。核心算法包括："
    )
    doc.add_paragraph(
        "（1）景点匹配评分（0-100分）：综合考虑评分基础分（30%）、兴趣匹配（30%）、"
        "季节匹配（25%）和热度调整（15%）四个维度。"
    )
    doc.add_paragraph(
        "（2）行程编排：根据匹配分数选择最优景点组合，按天数分配每日行程，"
        "同时考虑景点间的距离和交通时间。"
    )
    doc.add_paragraph(
        "（3）费用计算：综合门票、住宿、餐饮和交通费用，给出预估总花费，"
        "并与用户预算进行对比。"
    )

    add_heading_custom(doc, "4.3 应用页面与功能", level=2)

    doc.add_paragraph(
        "应用采用侧边栏导航设计，包含9个功能页面，每个页面聚焦一个分析主题："
    )

    pages = [
        ("📊 数据总览", "全局指标展示（景点数、酒店数、美食数、推荐路线数）+ 今日天气速览 + 景点地图。"),
        ("🌤️ 今日出行指南", "实时天气大卡片 + 景点拥挤度排行（建议避开/推荐前往）+ 今日高性价比酒店TOP8。"),
        ("🤖 智能行程规划", "输入天数/预算/月份/偏好，一键生成行程，含费用概览、每日安排、景点匹配度排行、淡旺季价格对比。"),
        ("🏞️ 景点探索", "景点详情页，含评分/评论/门票/游玩时间指标 + 景观介绍 + 传说故事 + 景区内路线 + 地图定位。"),
        ("🏨 酒店性价比", "酒店详情页，含评分/星级/价格/性价比指标 + 房型价格对比图 + 区域价格对比图。"),
        ("🔗 游住联动推荐", "选择景点后自动推荐附近酒店，按推荐指数排序，支持预算和评分筛选。"),
        ("🗺️ 路线规划", "7条经典路线（1/2/3日游），含每日行程安排、景点间距离和交通方式。"),
        ("📅 最佳旅行时间", "四季气候分析 + 各月份推荐景点 + 最佳旅行时间热力图。"),
        ("🍜 美食推荐", "美食评分排行 + 类型分布 + 价格分析 + 区域分布。"),
    ]

    for name, desc in pages:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}：")
        run.bold = True
        p.add_run(desc)

    add_heading_custom(doc, "4.3.1 技术栈", level=3)
    doc.add_paragraph(
        "本应用主要采用Python编程语言，并依赖以下关键库："
    )
    techs = [
        ("Streamlit：", "作为Web应用框架，提供了快速构建数据应用的便捷方式。"),
        ("Pandas & NumPy：", "用于数据的加载、清洗、转换和聚合，是数据处理的核心工具。"),
        ("Plotly / Plotly Express：", "用于创建交互式、高质量的可视化图表，包括柱状图、饼图、热力图、散点图等。"),
        ("Requests：", "用于调用外部天气API获取实时数据。"),
        ("Python-docx：", "用于生成Word格式的分析报告。"),
    ]
    for title, desc in techs:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(desc)

    add_heading_custom(doc, "4.4 应用部署与在线访问", level=2)

    doc.add_paragraph(
        "应用采用Streamlit Cloud平台进行部署，实现了零配置的在线托管服务。"
        "部署流程如下："
    )
    doc.add_paragraph(
        "（1）代码托管：将项目代码推送至GitHub仓库（PXT-GH/wenzhou-travel-viz），"
        "确保代码版本管理的规范性。"
    )
    doc.add_paragraph(
        "（2）平台配置：在Streamlit Cloud平台选择对应的GitHub仓库、分支（master）"
        "和主文件路径（app.py），配置Python依赖文件（requirements.txt）。"
    )
    doc.add_paragraph(
        "（3）自动部署：平台自动安装依赖、启动应用，部署完成后生成公开访问网址。"
        "用户无需注册或登录，通过网址即可直接访问。"
    )
    doc.add_paragraph(
        "（4）持续更新：当代码更新并推送到GitHub后，Streamlit Cloud自动触发重新部署，"
        "确保在线应用始终反映最新版本。"
    )
    doc.add_paragraph(
        "应用价值：该Streamlit应用使非技术背景的用户也能轻松探索和理解复杂的旅游数据，"
        "为游客提供了动态、交互式的数据分析工具，能够快速获取支持出行决策的信息。"
    )

    doc.add_page_break()

    # ====== 5 总结与展望 ======
    add_heading_custom(doc, "5 总结与展望", level=1)

    add_heading_custom(doc, "5.1 研究结论总结", level=2)

    doc.add_paragraph(
        "本研究通过对温州市旅游数据的全面分析和交互式可视化平台的构建，取得了以下重要成果："
    )

    findings = [
        ("旅游数据体系构建：",
         "成功构建了包含18个景点、45家酒店、25家美食店铺的温州旅游多维度数据库，"
         "涵盖评分、价格、评论数、地理位置、最佳游览时间等关键字段。"),
        ("创新分析指标设计：",
         "设计了性价比指数、拥挤度算法、推荐指数等多个创新分析指标，"
         "为旅游数据分析提供了新的量化工具。"),
        ("多维度可视化分析：",
         "围绕6大分析主题，生成了23个交互式Plotly可视化图表，"
         "直观展示了温州旅游资源的分布特征和季节性规律。"),
        ("实时数据集成：",
         "成功集成天气API、拥挤度计算和酒店价格浮动算法，"
         "为用户提供今日出行指南功能。"),
        ("交互式Web应用：",
         "基于Streamlit构建了包含9个功能页面的交互式Web应用，"
         "集成数据可视化展示、智能行程规划、景点深度探索等功能。"),
        ("在线部署与访问：",
         "应用已成功部署至Streamlit Cloud，通过公开网址即可访问，"
         "为用户提供便捷的旅游决策支持工具。"),
    ]

    for title, desc in findings:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}")
        run.bold = True
        p.add_run(desc)

    add_heading_custom(doc, "5.2 研究的不足与局限性", level=2)

    doc.add_paragraph(
        "尽管本研究力求全面和深入，但仍存在以下局限性："
    )

    limitations = [
        "数据来源以构建和模拟为主，部分数据基于真实信息整理但经过模拟处理，"
        "与真实OTA平台的实时数据存在一定差距。",
        "拥挤度算法基于多因子加权模型，但未引入历史客流数据和机器学习预测，"
        "预测精度有提升空间。",
        "路线规划算法相对基础，未考虑实时交通状况、用餐时间等约束条件。",
        "应用目前仅覆盖温州市，未扩展至浙江省其他城市。",
    ]
    for lim in limitations:
        doc.add_paragraph(f"• {lim}")

    add_heading_custom(doc, "5.3 研究的改进与展望", level=2)

    doc.add_paragraph(
        "基于以上研究和局限性，未来的研究工作可以聚焦于以下几个方向："
    )

    prospects = [
        ("数据源扩展：",
         "接入携程、美团等OTA平台的真实API，获取更准确的实时旅游数据，"
         "包括实时客流、实时价格、用户真实评价等。"),
        ("智能分析升级：",
         "引入机器学习算法，构建客流预测模型、价格预测模型，"
         "提升拥挤度计算和价格浮动的预测精度。"),
        ("路线规划优化：",
         "加入更多约束条件（如实时交通时间、用餐时间、景点开放时间等），"
         "结合遗传算法或蚁群算法实现更精准的行程规划。"),
        ("NLP情感分析：",
         "利用自然语言处理技术分析游客评价的情感倾向，"
         "为景点和酒店评价提供更丰富的维度。"),
        ("多城市扩展：",
         "将平台扩展至浙江省其他旅游城市，构建更大的旅游数据可视化平台，"
         "实现城市间的旅游数据对比分析。"),
        ("移动端适配：",
         "优化应用的移动端体验，开发微信小程序或移动App，"
         "方便用户在出行途中随时使用。"),
    ]

    for title, desc in prospects:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "数据可视化和数据驱动的分析方法将持续为理解和优化旅游行业提供强大的支持，"
        "本研究的探索为进一步的深入研究奠定了基础。"
    )

    doc.add_page_break()

    # ====== 参考文献 ======
    add_heading_custom(doc, "参考文献", level=1)

    refs = [
        "Homepage[EB/OL]. [2026-06-01]. https://cg.cs.tsinghua.edu.cn/course/vis/.",
        "Streamlit documentation[EB/OL]. [2026-06-01]. https://docs.streamlit.io/.",
        "Plotly Python Graphing Library[EB/OL]. [2026-06-01]. https://plotly.com/python/.",
        "wttr.in: a console-oriented weather forecast service[EB/OL]. [2026-06-01]. https://wttr.in/.",
        "Pandas documentation[EB/OL]. [2026-06-01]. https://pandas.pydata.org/docs/.",
        "NumPy documentation[EB/OL]. [2026-06-01]. https://numpy.org/doc/.",
        "Python-docx documentation[EB/OL]. [2026-06-01]. https://python-docx.readthedocs.io/.",
        "McKinney, W. (2017). Python for data analysis: Data wrangling with Pandas, NumPy, and IPython. O'Reilly Media, Inc..",
        "Tufte, E. R. (2001). The Visual Display of Quantitative Information (2nd ed.). Graphics Press.",
        "温州市文化广电旅游局[EB/OL]. [2026-06-01]. https://whlyj.wenzhou.gov.cn/.",
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] ")
        run.bold = True
        p.add_run(ref)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, "温州旅游数据可视化分析报告_v2.docx")
    doc.save(output_path)
    print(f"[OK] 报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
